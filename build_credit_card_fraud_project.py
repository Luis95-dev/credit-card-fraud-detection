from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("Credit_Card_Fraud_Detection_Final_Project.docx")

NAVY = "203748"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MID_BLUE = "2B5163"
BODY = "202124"
MUTED = "5F6B76"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BORDER = "B8C4CE"


def set_cell_text(cell, text, bold=False, color=BODY, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run_font(run, "Calibri", size, color, bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_left_border(paragraph, color=BLUE, size=18, space=10):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=9, color=MUTED)


def add_numbering_definition(doc, kind):
    numbering = doc.part.numbering_part.element
    existing_abstract = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    existing_nums = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(existing_abstract, default=-1) + 1
    num_id = max(existing_nums, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if kind == "decimal" else "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if kind == "decimal" else "\u2022")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Symbol")
        fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_body_paragraph(doc, text, keep_with_next=False):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    return paragraph


def add_bullet(doc, num_id, label, text):
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, num_id)
    label_run = paragraph.add_run(label)
    set_run_font(label_run, bold=True, color=DARK_BLUE)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, color=BODY)
    return paragraph


def add_numbered_step(doc, num_id, label, text):
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, num_id)
    paragraph.paragraph_format.keep_with_next = False
    label_run = paragraph.add_run(label)
    set_run_font(label_run, bold=True, color=DARK_BLUE)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, color=BODY)
    return paragraph


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BODY)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(29)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = RGBColor.from_string(MID_BLUE)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(4)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def add_model_table(doc):
    rows = [
        ("Logistic Regression", "Linear baseline", "Fast, transparent, and produces interpretable probability estimates.", "May miss complex relationships unless features are carefully prepared."),
        ("Decision Tree", "Rule-based baseline", "Easy to explain and can model nonlinear patterns.", "Can overfit and may change noticeably with small data differences."),
        ("Random Forest", "Ensemble model", "Usually more stable than one tree and can capture complex interactions.", "Less transparent and more computationally demanding."),
    ]
    table = doc.add_table(rows=1, cols=4)
    headers = ["Model", "Project role", "Primary strength", "Main limitation"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE, size=9.5)
        set_cell_shading(table.rows[0].cells[idx], NAVY)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            set_cell_text(cells[col_idx], value, size=9.25)
            if row_idx % 2 == 1:
                set_cell_shading(cells[col_idx], LIGHT_GRAY)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [1900, 1800, 3000, 2660])
    set_table_cell_margins(table)
    set_table_borders(table)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    return table


def add_metrics_table(doc):
    rows = [
        ("Precision", "Of the transactions flagged as fraud, how many are truly fraudulent?", "Shows the false-alarm burden placed on analysts and customers."),
        ("Recall", "Of all actual fraudulent transactions, how many are detected?", "Measures how much fraud the model successfully captures."),
        ("F1-score", "What is the balance between precision and recall?", "Provides one summary score when both missed fraud and false alerts matter."),
        ("ROC-AUC", "How well does the model rank fraud above normal activity across thresholds?", "Useful for broad model comparison, but it can appear optimistic with severe imbalance."),
        ("Precision-Recall AUC", "How strong is the precision-recall tradeoff across thresholds?", "Focuses on the minority class and is especially informative for rare fraud cases."),
    ]
    table = doc.add_table(rows=1, cols=3)
    headers = ["Metric", "Question answered", "Why it matters"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE, size=9.5)
        set_cell_shading(table.rows[0].cells[idx], NAVY)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            set_cell_text(cells[col_idx], value, size=9.2)
            if row_idx % 2 == 1:
                set_cell_shading(cells[col_idx], LIGHT_GRAY)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [1700, 3500, 4160])
    set_table_cell_margins(table)
    set_table_borders(table)
    return table


def build_document():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = ""
    header_p.paragraph_format.space_after = Pt(0)
    header_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    header_left = header_p.add_run("CREDIT CARD FRAUD DETECTION")
    set_run_font(header_left, size=8.5, color=MUTED, bold=True)
    header_right = header_p.add_run("\tFINAL PROJECT")
    set_run_font(header_right, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    page_text = footer_p.add_run("Page ")
    set_run_font(page_text, size=9, color=MUTED)
    add_page_field(footer_p)

    # Cover page: editorial-cover pattern with restrained academic styling.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(104)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    kicker_run = kicker.add_run("FINAL PROJECT | MACHINE LEARNING")
    set_run_font(kicker_run, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph(style="Title")
    title.add_run("Credit Card Fraud Detection Using Machine Learning and Imbalanced Data Handling")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("An Academic Project Proposal")

    topics = doc.add_paragraph()
    topics.alignment = WD_ALIGN_PARAGRAPH.CENTER
    topics.paragraph_format.space_before = Pt(18)
    topics.paragraph_format.space_after = Pt(0)
    topics_run = topics.add_run("Classification | Imbalanced Learning | Fraud Analytics")
    set_run_font(topics_run, size=10.5, color=MUTED, italic=True)

    cover_space = doc.add_paragraph()
    cover_space.paragraph_format.space_after = Pt(118)
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_after = Pt(0)
    date_run = date_p.add_run("College Class Final Project\nJune 2026")
    set_run_font(date_run, size=10.5, color=MUTED)
    date_p.add_run().add_break(WD_BREAK.PAGE)

    # Abstract
    doc.add_heading("Abstract", level=1)
    abstract = (
        "Credit card fraud creates financial losses, disrupts customer trust, and increases the cost of reviewing "
        "suspicious transactions. Detecting fraud quickly is therefore an important goal for banks, payment processors, "
        "and other financial institutions. This project proposes a machine learning approach to classify credit card "
        "transactions as legitimate or fraudulent. The central challenge is class imbalance: fraudulent transactions "
        "represent only a very small portion of the available data, while normal transactions form the overwhelming "
        "majority. As a result, a model may report very high accuracy simply by predicting almost every transaction as "
        "normal, even though it fails to detect the cases that matter most. The project will compare Logistic Regression, "
        "Decision Tree, and Random Forest models under both untreated and imbalance-aware training conditions. Techniques "
        "such as the Synthetic Minority Over-sampling Technique (SMOTE), random undersampling, and class weights will be "
        "evaluated to determine whether they improve minority-class detection without creating an unacceptable number of "
        "false alerts. Model performance will be measured using Precision, Recall, F1-score, ROC-AUC, and Precision-Recall "
        "AUC, supported by confusion matrices and threshold analysis. The expected outcome is a clear comparison of model "
        "and sampling choices, along with practical guidance on selecting a fraud detection approach. This work may help "
        "financial institutions reduce missed fraud, control investigation costs, and make more informed decisions when "
        "deploying machine learning systems."
    )
    abstract_p = add_body_paragraph(doc, abstract)
    abstract_p.paragraph_format.first_line_indent = Inches(0.3)

    # Introduction
    doc.add_heading("Introduction", level=1)
    add_body_paragraph(
        doc,
        "The continued growth of online shopping, mobile banking, and digital payment services has made credit card "
        "transactions faster and more convenient. At the same time, these systems create opportunities for fraudulent "
        "activity. Traditional rule-based controls remain useful, but fixed rules may struggle when fraud patterns change "
        "or when criminals imitate normal customer behavior. Machine learning can support fraud prevention by learning "
        "patterns from historical transactions and assigning a risk score to new activity."
    )
    add_body_paragraph(
        doc,
        "Fraud detection is not a standard classification problem because the two classes are extremely unequal. In a "
        "typical transaction dataset, legitimate purchases greatly outnumber fraudulent ones. A learning algorithm may "
        "therefore favor the majority class unless the training process explicitly addresses imbalance. The project will "
        "study this issue directly rather than treating it as a minor preprocessing concern."
    )

    # Problem Statement
    doc.add_heading("Problem Statement", level=1)
    add_body_paragraph(
        doc,
        "The project must identify rare fraudulent transactions while limiting unnecessary alerts on legitimate customer "
        "activity. These goals compete with each other: increasing sensitivity to fraud may also increase false positives, "
        "while reducing false positives may allow more fraudulent transactions to pass undetected."
    )
    callout = doc.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.16)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.line_spacing = 1.18
    set_paragraph_shading(callout, LIGHT_BLUE)
    set_paragraph_left_border(callout)
    key_run = callout.add_run("Why accuracy is not enough. ")
    set_run_font(key_run, bold=True, color=DARK_BLUE)
    callout_run = callout.add_run(
        "If 99.8% of transactions are legitimate, a model that labels every transaction as legitimate will achieve "
        "99.8% accuracy while detecting 0% of the fraud. Accuracy hides this failure because it is dominated by the "
        "majority class."
    )
    set_run_font(callout_run, color=BODY)
    add_body_paragraph(
        doc,
        "A useful solution must therefore be judged by its treatment of the minority class. Precision shows whether fraud "
        "alerts are trustworthy, recall shows whether actual fraud is captured, and the F1-score summarizes their balance. "
        "ROC-AUC and Precision-Recall AUC provide threshold-independent comparisons, with Precision-Recall AUC receiving "
        "special attention because it is more informative when positive cases are rare."
    )

    # Objective
    doc.add_heading("Project Objective", level=1)
    add_body_paragraph(
        doc,
        "The main objective is to design and evaluate an understandable machine learning workflow for credit card fraud "
        "detection that gives proper attention to class imbalance. The project will:"
    )
    bullet_num_id = add_numbering_definition(doc, "bullet")
    add_bullet(doc, bullet_num_id, "Compare baseline models. ", "Train Logistic Regression, Decision Tree, and Random Forest classifiers under consistent data splits.")
    add_bullet(doc, bullet_num_id, "Test imbalance treatments. ", "Evaluate SMOTE, random undersampling, and class-weighted learning against untreated baselines.")
    add_bullet(doc, bullet_num_id, "Use suitable metrics. ", "Compare precision, recall, F1-score, ROC-AUC, and Precision-Recall AUC rather than relying on accuracy alone.")
    add_bullet(doc, bullet_num_id, "Examine operational tradeoffs. ", "Study how decision thresholds affect missed fraud and false alerts.")
    add_bullet(doc, bullet_num_id, "Produce practical guidance. ", "Recommend a model and imbalance strategy that financial institutions could evaluate further in a real operating environment.")

    # Methodology
    doc.add_heading("Proposed Methodology", level=1)
    add_body_paragraph(
        doc,
        "The study will use a labeled credit card transaction dataset containing legitimate and fraudulent examples. "
        "Personally identifying information will not be required. The workflow will preserve a separate test set so that "
        "the final evaluation reflects previously unseen data."
    )
    decimal_num_id = add_numbering_definition(doc, "decimal")
    add_numbered_step(doc, decimal_num_id, "Data preparation and exploration. ", "Inspect class distribution, missing values, feature ranges, and transaction patterns. Divide the data using stratified training, validation, and test splits so each partition maintains a realistic fraud proportion.")
    add_numbered_step(doc, decimal_num_id, "Baseline modeling. ", "Train Logistic Regression, Decision Tree, and Random Forest models without resampling. Standardize numeric features when needed for Logistic Regression, while allowing tree-based models to use the original feature scales.")
    add_numbered_step(doc, decimal_num_id, "Imbalance handling. ", "Create separate experiments using class weights, SMOTE, and random undersampling. Apply resampling only to the training data or training fold to prevent information from the validation or test set from leaking into model development.")
    add_numbered_step(doc, decimal_num_id, "Model validation. ", "Use stratified cross-validation and limited hyperparameter tuning to compare models fairly. Keep the same folds and evaluation rules across experiments.")
    add_numbered_step(doc, decimal_num_id, "Performance evaluation. ", "Report confusion matrices, precision, recall, F1-score, ROC-AUC, and Precision-Recall AUC. Review multiple probability thresholds rather than assuming that the default 0.50 cutoff is appropriate for fraud detection.")
    add_numbered_step(doc, decimal_num_id, "Final comparison. ", "Select the strongest approach based on minority-class performance, stability, interpretability, and the likely cost of false negatives and false positives.")

    doc.add_heading("Models Included in the Comparison", level=2)
    add_model_table(doc)

    doc.add_heading("Evaluation Framework", level=2)
    add_body_paragraph(
        doc,
        "The final comparison will emphasize the following metrics. Accuracy will still be reported for completeness, but "
        "it will not be used as the primary decision criterion."
    )
    add_metrics_table(doc)

    # Expected Results
    doc.add_heading("Expected Results", level=1)
    add_body_paragraph(
        doc,
        "All three models are expected to produce high overall accuracy because most transactions are legitimate. However, "
        "the untreated models may show weaker recall and Precision-Recall AUC, revealing that accuracy alone overstates "
        "their usefulness. Class weights, SMOTE, or undersampling are expected to improve the detection of fraudulent "
        "transactions, although the improvement may be accompanied by more false positives."
    )
    add_body_paragraph(
        doc,
        "Random Forest may provide the strongest overall ranking performance because it can represent nonlinear patterns "
        "and interactions. Logistic Regression may remain competitive and offer clearer interpretation, while a single "
        "Decision Tree may be easy to explain but less stable. These are expectations rather than guaranteed outcomes; "
        "the project will base its conclusions on the measured results."
    )
    add_body_paragraph(
        doc,
        "The most valuable result will be a transparent explanation of the tradeoff between catching more fraud and "
        "creating more alerts. Financial institutions could use this comparison to choose a model, resampling strategy, "
        "and threshold that match their fraud losses, review capacity, customer experience goals, and regulatory duties."
    )

    # Conclusion
    doc.add_heading("Conclusion", level=1)
    add_body_paragraph(
        doc,
        "Credit card fraud detection requires more than selecting a classifier with high accuracy. Because fraudulent "
        "transactions are rare, the project must focus on how well a model recognizes the minority class and how many "
        "false alerts it creates. By comparing Logistic Regression, Decision Tree, and Random Forest models with SMOTE, "
        "undersampling, and class weights, this project will demonstrate how imbalance handling changes model behavior."
    )
    add_body_paragraph(
        doc,
        "The project will also recognize the limits of a classroom study. A historical dataset cannot represent every "
        "future fraud pattern, and the practical cost of a false negative or false positive may differ across financial "
        "institutions. Data privacy, model explainability, changing customer behavior, and continuous monitoring would all "
        "require additional attention before a model could be used in a live payment system. The findings will therefore "
        "be presented as comparative evidence and a foundation for further testing rather than as a production-ready "
        "solution."
    )
    add_body_paragraph(
        doc,
        "Using precision, recall, F1-score, ROC-AUC, and Precision-Recall AUC will support a fair and practical evaluation. "
        "The final outcome will be an accessible machine learning study that connects technical model performance to the "
        "real needs of financial institutions: reducing fraud losses, protecting customers, and using investigation "
        "resources responsibly."
    )

    # Document metadata.
    properties = doc.core_properties
    properties.title = "Credit Card Fraud Detection Using Machine Learning and Imbalanced Data Handling"
    properties.subject = "College class final project on machine learning and imbalanced fraud data"
    properties.author = "Student"
    properties.keywords = "credit card fraud, machine learning, imbalanced data, SMOTE, class weights"

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(path.resolve())
